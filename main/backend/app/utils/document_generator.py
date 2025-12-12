"""
Document Generation Utilities
Generates PDF and DOCX documents from templates with proper formatting
"""
import io
import re
from pathlib import Path
from typing import Dict, Any, Tuple, Optional
from datetime import datetime
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocumentGenerator:
    """Generate professional PDF and Word documents from templates"""
    
    # Deprecated default company config; header now uses user-provided data only
    COMPANY_CONFIG = {}
    
    # Font configuration - using Courier (Delivery-like monospace)
    FONTS = {
        "title": ("Helvetica-Bold", 16),
        "heading": ("Helvetica-Bold", 12),
        "body": ("Courier", 11),  # Delivery-like font
        "small": ("Courier", 9)
    }
    
    @classmethod
    def generate_pdf_from_template(
        cls,
        template_path: Optional[Path],
        data: Dict[str, Any],
        output_stream: io.BytesIO
    ) -> io.BytesIO:
        """
        Generate PDF from template by replacing text placeholders
        Uses Courier font (similar to Delivery)
        """
        if template_path:
            # Read template PDF
            reader = PdfReader(str(template_path))
            template_text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    template_text += page_text + "\n"
        else:
            # Use content directly from data
            template_text = data.get('content', '')
        
        # Add company details to data if not provided
        enhanced_data = {**data}
        
        # DO NOT use defaults - only use user-provided data
        # If user didn't provide company info, leave placeholders empty
        print(f"🔍 User-provided company_name: {enhanced_data.get('company_name')}")
        print(f"🔍 User-provided company_address: {enhanced_data.get('company_address')}")
        print(f"🔍 User-provided contact_info: {enhanced_data.get('contact_info')}")
        
        # Debug: log template text and data
        print(f"🔍 PDF GENERATION DEBUG:")
        print(f"Template text (first 500 chars): {template_text[:500]}")
        print(f"Enhanced data keys: {list(enhanced_data.keys())}")
        print(f"Enhanced data values: {enhanced_data}")
        
        # Replace all placeholders in template text
        content = cls._replace_placeholders(template_text, enhanced_data)
        
        print(f"After replacement - checking for remaining placeholders:")
        import re
        remaining_placeholders = re.findall(r'\[([^\]]+)\]', content)
        if remaining_placeholders:
            print(f"⚠️ WARNING: These placeholders were NOT replaced: {remaining_placeholders}")
        print(f"After replacement (first 500 chars): {content[:500]}")
        
        # Create new PDF with filled content
        doc = SimpleDocTemplate(
            output_stream,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=0.75*inch
        )
        
        # Build story with proper formatting
        story = []
        
        # Add company header
        story.extend(cls._create_pdf_header(enhanced_data))
        story.append(Spacer(1, 0.3*inch))
        
        # Add content with preserved formatting
        story.extend(cls._create_pdf_content(content, enhanced_data))
        
        # Build PDF
        doc.build(story)
        output_stream.seek(0)
        return output_stream
    
    @classmethod
    def generate_docx_from_template(
        cls,
        template_path: Optional[Path],
        data: Dict[str, Any],
        output_stream: io.BytesIO
    ) -> io.BytesIO:
        """Generate Word document from template"""
        if template_path:
            # Read template PDF to extract text
            reader = PdfReader(str(template_path))
            template_text = ""
            for page in reader.pages:
                template_text += page.extract_text() + "\n"
        else:
            # Use content directly from data
            template_text = data.get('content', '')
        
        # Replace placeholders
        content = cls._replace_placeholders(template_text, data)
        
        # Create Word document
        doc = Document()
        
        # Set default font to Courier New (Delivery-like)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(11)
        
        # Add company header
        cls._add_docx_header(doc, data)
        
        # Add content
        cls._add_docx_content(doc, content)
        
        # Save to stream
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream
    
    @classmethod
    def _replace_placeholders(cls, template_text: str, data: Dict[str, Any]) -> str:
        """Replace [PLACEHOLDER] with actual values - handles multiple formats"""
        content = template_text
        
        print(f"🔍 _replace_placeholders called with data keys: {list(data.keys())}")
        
        # Replace placeholders for each field in data
        for key, value in data.items():
            if value is None:
                value = ""
            
            # Convert value to string
            value_str = str(value)
            
            # Treat 'skip' for signatory fields as empty string (to avoid rendering 'skip')
            if key in ("signatory_name", "signatory_designation"):
                if value_str.strip().lower() == "skip":
                    value_str = ""
            
            print(f"🔍 Processing field '{key}' with value '{value_str}'")
            
            # Try multiple placeholder formats:
            # 1. Exact match: [employee_name]
            placeholder = f"[{key}]"
            content = content.replace(placeholder, value_str)
            
            # 2. Uppercase: [EMPLOYEE_NAME]
            placeholder_upper = f"[{key.upper()}]"
            content = content.replace(placeholder_upper, value_str)
            
            # 3. Title case: [Employee_Name]
            placeholder_title = f"[{key.title()}]"
            content = content.replace(placeholder_title, value_str)
            
            # 4. With spaces instead of underscores: [employee name]
            key_spaced = key.replace('_', ' ')
            placeholder_spaced = f"[{key_spaced}]"
            content = content.replace(placeholder_spaced, value_str)
            
            # 5. Spaces + uppercase: [EMPLOYEE NAME]
            placeholder_spaced_upper = f"[{key_spaced.upper()}]"
            content = content.replace(placeholder_spaced_upper, value_str)
            
            # 6. Spaces + title case: [Employee Name]
            placeholder_spaced_title = f"[{key_spaced.title()}]"
            content = content.replace(placeholder_spaced_title, value_str)
            
            # 7. Handle common variations
            if key == "employee_name" or key == "candidate_name":
                content = content.replace("[Name]", value_str)
                content = content.replace("[NAME]", value_str)
                content = content.replace("[Employee Name]", value_str)
                content = content.replace("[EMPLOYEE NAME]", value_str)
            elif key == "designation" or key == "position":
                content = content.replace("[Position]", value_str)
                content = content.replace("[POSITION]", value_str)
                content = content.replace("[Job Title]", value_str)
                content = content.replace("[Designation]", value_str)
                content = content.replace("[DESIGNATION]", value_str)
            elif key == "joining_date":
                print(f"🔍 Replacing [Date of Joining] with '{value_str}'")
                content = content.replace("[Start Date]", value_str)
                content = content.replace("[Date of Joining]", value_str)
                content = content.replace("[DATE OF JOINING]", value_str)
            elif key == "salary":
                print(f"🔍 Replacing [CTC] with '{value_str}'")
                content = content.replace("[CTC]", value_str)
                content = content.replace("[Annual CTC]", value_str)
                content = content.replace("[ANNUAL CTC]", value_str)
            elif key == "last_working_date":
                content = content.replace("[Relieving Date]", value_str)
                content = content.replace("[RELIEVING DATE]", value_str)
            elif key == "offer_acceptance_date":
                content = content.replace("[Offer Acceptance Date]", value_str)
                content = content.replace("[OFFER ACCEPTANCE DATE]", value_str)
            elif key == "confirmation_date":
                content = content.replace("[Confirmation Date]", value_str)
                content = content.replace("[CONFIRMATION DATE]", value_str)
            elif key == "bonus_amount":
                content = content.replace("[Bonus Amount]", value_str)
                content = content.replace("[BONUS AMOUNT]", value_str)
            elif key == "reason_for_termination":
                content = content.replace("[Reason for Termination]", value_str)
                content = content.replace("[REASON FOR TERMINATION]", value_str)
            elif key == "termination_date":
                content = content.replace("[Termination Date]", value_str)
                content = content.replace("[TERMINATION DATE]", value_str)
            elif key == "company_name":
                content = content.replace("[Company Name]", value_str)
                content = content.replace("[COMPANY NAME]", value_str)
            elif key == "company_address":
                content = content.replace("[Company Address]", value_str)
                content = content.replace("[COMPANY ADDRESS]", value_str)
            elif key == "contact_info":
                content = content.replace("[Contact Info]", value_str)
                content = content.replace("[CONTACT INFO]", value_str)
            elif key == "signatory_name":
                content = content.replace("[Signatory Name]", value_str)
                content = content.replace("[SIGNATORY NAME]", value_str)
            elif key == "signatory_designation":
                content = content.replace("[Signatory Designation]", value_str)
                content = content.replace("[SIGNATORY DESIGNATION]", value_str)
            elif key == "pronoun_subject":
                content = content.replace("[he/she/they]", value_str)
            elif key == "pronoun_possessive":
                content = content.replace("[his/her/their]", value_str)
        
        return content
    
    @classmethod
    def _create_pdf_header(cls, data: Dict[str, Any]) -> list:
        """Create PDF header with company info from data - NO DEFAULTS"""
        styles = getSampleStyleSheet()
        story = []
        
        # Get company info from user data ONLY - no defaults
        company_name = data.get("company_name", "")
        company_address = data.get("company_address", "")
        contact_info = data.get("contact_info", "")
        
        # Only add if user provided
        if company_name and company_name != "skip":
            company_style = ParagraphStyle(
                'CompanyName',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000000'),
                spaceAfter=6,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            story.append(Paragraph(company_name, company_style))
        
        # Address - only if provided
        if company_address and company_address != "skip":
            address_style = ParagraphStyle(
                'Address',
                parent=styles['Normal'],
                fontSize=9,
                alignment=TA_CENTER,
                fontName='Courier'
            )
            story.append(Paragraph(company_address, address_style))
            
            # Contact info - only if provided  
            if contact_info and contact_info != "skip":
                story.append(Paragraph(contact_info, address_style))
        
        if story:  # Only add spacer if we added header
            story.append(Spacer(1, 0.3*inch))
        
        return story
        
        return story
    
    @classmethod
    def _create_pdf_content(cls, content: str, data: Dict[str, Any] = None) -> list:
        """Create PDF content paragraphs with signature support"""
        styles = getSampleStyleSheet()
        story = []
        
        # Body style with Courier font
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Courier',
            leading=14,
            spaceAfter=12
        )
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        # Check if signature should replace name+designation
        has_signature = bool(data and data.get('signatory_signature'))
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # If we detect signatory placeholders (rare after replacement) and have signature, embed image
            if has_signature and ('[Signatory Name]' in para_text or '[SIGNATORY NAME]' in para_text or 
                                   '[Signatory Designation]' in para_text or '[SIGNATORY DESIGNATION]' in para_text):
                # Embed signature image instead of text
                try:
                    import base64
                    import tempfile
                    from reportlab.platypus import Image 
                    
                    signature_base64 = data["signatory_signature"]
                    # Remove data URL prefix if present
                    if ',' in signature_base64:
                        signature_base64 = signature_base64.split(',')[1]
                    
                    signature_bytes = base64.b64decode(signature_base64)
                    
                    # Create temporary image from bytes
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                        tmp.write(signature_bytes)
                        tmp_path = tmp.name
                    
                    # Add signature image to PDF (larger size since it contains name+designation)
                    story.append(Spacer(1, 0.3*inch))
                    sig_img = Image(tmp_path, width=3*inch, height=1.5*inch)
                    story.append(sig_img)
                    
                    # Clean up temp file
                    import os
                    os.unlink(tmp_path)
                    
                    print("✅ Embedded signature image (replaces name+designation)")
                    continue  # Skip adding text paragraph
                    
                except Exception as e:
                    print(f"⚠️ Error embedding signature: {e}")
                    # Fall through to add text paragraph as fallback
            
            # Check if it's a title/heading
            if para_text.isupper() or para_text.endswith('Letter'):
                heading_style = ParagraphStyle(
                    'Heading',
                    parent=styles['Heading2'],
                    fontSize=14,
                    fontName='Helvetica-Bold',
                    alignment=TA_CENTER,
                    spaceAfter=12
                )
                story.append(Paragraph(para_text, heading_style))
            else:
                # Regular paragraph
                para_text = para_text.replace('\n', '<br/>')
                story.append(Paragraph(para_text, body_style))

                # If this is the sign-off line, place signature right after
                if has_signature and para_text.lower().startswith('sincerely'):
                    try:
                        import base64
                        from reportlab.platypus import Image
                        signature_base64 = data["signatory_signature"]
                        if ',' in signature_base64:
                            signature_base64 = signature_base64.split(',')[1]
                        signature_bytes = base64.b64decode(signature_base64)
                        import tempfile
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                            tmp.write(signature_bytes)
                            tmp_path = tmp.name
                        story.append(Spacer(1, 0.2*inch))
                        sig_img = Image(tmp_path, width=2.5*inch, height=1.2*inch)
                        story.append(sig_img)
                        import os
                        os.unlink(tmp_path)
                    except Exception as e:
                        print(f"⚠️ Error embedding signature after sign-off: {e}")
        
        # Fallback: add signature image near the end if not already added
        if has_signature:
            try:
                import base64
                from reportlab.platypus import Image
                signature_base64 = data["signatory_signature"]
                if ',' in signature_base64:
                    signature_base64 = signature_base64.split(',')[1]
                signature_bytes = base64.b64decode(signature_base64)
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    tmp.write(signature_bytes)
                    tmp_path = tmp.name
                story.append(Spacer(1, 0.2*inch))
                sig_img = Image(tmp_path, width=2.5*inch, height=1.2*inch)
                story.append(sig_img)
                import os
                os.unlink(tmp_path)
            except Exception as e:
                print(f"⚠️ Error embedding signature (fallback): {e}")
        
        return story
    
    @classmethod
    def _add_docx_header(cls, doc: Document, data: Dict[str, Any]):
        """Add company header to Word document - uses ONLY user-provided data"""
        # Get company info from user data ONLY
        company_name = data.get("company_name", "")
        company_address = data.get("company_address", "")
        contact_info = data.get("contact_info", "")
        
        # Only add header elements if user provided the data
        if company_name:
            heading = doc.add_heading(company_name, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if company_address:
            address_para = doc.add_paragraph(company_address)
            address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            address_para.runs[0].font.size = Pt(9)
            address_para.runs[0].font.name = 'Courier New'
        
        if contact_info:
            contact_para = doc.add_paragraph(contact_info)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.runs[0].font.size = Pt(9)
            contact_para.runs[0].font.name = 'Courier New'
        
        # Spacing
        doc.add_paragraph()
    
    @classmethod
    def _add_docx_content(cls, doc: Document, content: str):
        """Add content to Word document"""
        # Split into paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Check if heading
            if para_text.isupper() or para_text.endswith('Letter'):
                heading = doc.add_heading(para_text, level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Regular paragraph
                para = doc.add_paragraph(para_text)
                para.runs[0].font.name = 'Courier New'
                para.runs[0].font.size = Pt(11)
    
    @classmethod
    def generate_preview_html(
        cls,
        template_path: Path,
        data: Dict[str, Any]
    ) -> str:
        """Generate HTML preview of the document"""
        # Read template
        reader = PdfReader(str(template_path))
        template_text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                template_text += page_text + "\n"
        
        # Replace placeholders (with 'skip' treated as empty for signatory fields)
        content = cls._replace_placeholders(template_text, data)
        
        # Company header values from user-provided data only
        company_name = data.get("company_name", "")
        company_address = data.get("company_address", "")
        contact_info = data.get("contact_info", "")
        
        # Convert to HTML
        # Build signature HTML if provided
        signature_html = ""
        signature_b64 = data.get("signatory_signature", "")
        if signature_b64:
            # If signature contains a data URL, use as-is; otherwise prefix PNG
            if signature_b64.startswith("data:image"):
                signature_src = signature_b64
            else:
                signature_src = f"data:image/png;base64,{signature_b64}"
            signature_html = f"<div style=\"margin:12px 0;\"><img alt=\"Signature\" src=\"{signature_src}\" style=\"height:96px;\" /></div>"

        html = f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{
            font-family: 'Courier New', Courier, monospace;
            max-width: 800px;
            margin: 40px auto;
            padding: 40px;
            background: white;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
            line-height: 1.6;
        }}
        .header {{
            text-align: center;
            border-bottom: 2px solid #333;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        .company-name {{
            font-size: 24px;
            font-weight: bold;
            margin-bottom: 10px;
        }}
        .company-info {{
            font-size: 11px;
            color: #666;
        }}
        .content {{
            font-size: 11pt;
            white-space: pre-wrap;
        }}
        .letter-title {{
            text-align: center;
            font-size: 18px;
            font-weight: bold;
            margin: 30px 0;
        }}
    </style>
</head>
<body>
    <div class="header">
        {f'<div class="company-name">{company_name}</div>' if company_name and company_name != 'skip' else ''}
        {f'<div class="company-info">{company_address}</div>' if company_address and company_address != 'skip' else ''}
        {f'<div class="company-info">{contact_info}</div>' if contact_info and contact_info != 'skip' else ''}
    </div>
    <div class="content">{cls._format_html_content(content)}</div>
</body>
</html>
"""
        # Inject signature inline by replacing signatory placeholders if present
        if signature_html:
            # Replace both title-case and uppercase placeholders inside the HTML-formatted content
            html = html.replace("[Signatory Name]", signature_html)
            html = html.replace("[SIGNATORY NAME]", signature_html)
            html = html.replace("[Signatory Designation]", "")
            html = html.replace("[SIGNATORY DESIGNATION]", "")
            # Fallback: if no placeholders were present, append signature near the end
            if signature_html not in html:
                html = html.replace("</div>\n </body>", f"{signature_html}</div>\n </body>")
        return html
    
    @classmethod
    def _format_html_content(cls, content: str) -> str:
        """Format content for HTML display"""
        # Escape HTML
        content = content.replace('<', '&lt;').replace('>', '&gt;')
        
        # Convert line breaks
        content = content.replace('\n\n', '</p><p>')
        content = content.replace('\n', '<br/>')
        
        # Wrap in paragraphs
        content = f'<p>{content}</p>'
        
        return content
