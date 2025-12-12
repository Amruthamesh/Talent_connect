"""
Metadata extraction utility for various document types.
Handles both selectable text documents and non-selectable text documents (images, scanned PDFs).
"""
import re
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
from PyPDF2 import PdfReader
from docx import Document
import json


class DocumentMetadata:
    """Represents extracted metadata from a document."""
    
    def __init__(self):
        self.file_name: str = ""
        self.file_type: str = ""
        self.file_size: int = 0
        self.file_path: str = ""
        self.extracted_at: str = datetime.now().isoformat()
        self.content_type: str = ""  # e.g., "selectable_text", "scanned_document", "mixed"
        self.page_count: Optional[int] = None
        self.language: str = "en"
        self.has_selectable_text: bool = True
        self.has_images: bool = False
        self.image_count: int = 0
        self.is_scanned: bool = False
        self.ocr_confidence: Optional[float] = None  # Confidence level if OCR was used
        self.text_length: int = 0
        self.sections: Dict[str, Any] = {}  # Detected sections like "experience", "education", etc.
        self.entities: Dict[str, List[str]] = {  # Extracted entities
            "emails": [],
            "phones": [],
            "urls": [],
            "dates": [],
            "locations": [],
            "skills": [],
            "degrees": [],
            "companies": [],
            "certifications": [],
        }
        self.metadata_tags: Dict[str, str] = {}  # General metadata key-value pairs


class MetadataExtractor:
    """Extracts comprehensive metadata from various document types."""
    
    # Common skill keywords (can be expanded)
    SKILL_KEYWORDS = {
        # Programming Languages
        "python", "java", "javascript", "typescript", "c++", "c#", "golang", "rust", "php", "ruby", "swift",
        "kotlin", "scala", "r", "matlab", "perl", "groovy",
        # Web Frameworks
        "react", "vue", "angular", "fastapi", "django", "flask", "spring", "express", "nodejs", "node.js",
        # Databases
        "sql", "mysql", "postgresql", "mongodb", "cassandra", "redis", "elasticsearch", "oracle", "sqlite",
        # Cloud & DevOps
        "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "gitlab", "github", "terraform",
        # Data & ML
        "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "spark", "hadoop", "kafka",
        # Other Tech
        "git", "linux", "windows", "macos", "rest api", "graphql", "html", "css", "json", "xml",
        "agile", "scrum", "jira", "confluence", "slack", "salesforce",
    }
    
    # Common degree patterns
    DEGREE_PATTERNS = [
        r"(?:bachelor|b\.?s\.?|b\.?a\.?|b\.?sc\.?|b\.?tech\.?)",
        r"(?:master|m\.?s\.?|m\.?a\.?|m\.?sc\.?|m\.?tech\.?|m\.?b\.?a\.?)",
        r"(?:phd|ph\.?d\.?|doctorate)",
        r"(?:diploma|associate)",
        r"(?:certification|certified)",
    ]
    
    # Company keywords (common patterns)
    COMPANY_PATTERNS = [
        r"(?:worked|worked at|employed|employed at|at)\s+([A-Z][A-Za-z\s&.,\-]+)",
        r"([A-Z][A-Za-z\s&.,\-]+)\s+(?:inc|corp|ltd|llc|gmbh|ag|sa)",
    ]
    
    @staticmethod
    def _get_file_metadata(path: Path) -> Dict[str, Any]:
        """Extract basic file metadata."""
        stat = path.stat()
        return {
            "file_name": path.name,
            "file_type": path.suffix.lower(),
            "file_size": stat.st_size,
            "file_path": str(path),
        }
    
    @staticmethod
    def _extract_text_from_pdf(path: Path) -> tuple[str, bool, int]:
        """
        Extract text from PDF.
        Returns: (text, has_selectable_text, page_count)
        """
        try:
            reader = PdfReader(path)
            page_count = len(reader.pages)
            pages = []
            has_selectable = False
            
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    has_selectable = True
                pages.append(page_text)
            
            text = "\n".join(pages)
            return text, has_selectable, page_count
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def _extract_text_from_docx(path: Path) -> tuple[str, int]:
        """
        Extract text from DOCX.
        Returns: (text, paragraph_count)
        """
        try:
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            return text, len(doc.paragraphs)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def _extract_text_from_plain_text(path: Path) -> str:
        """Extract text from plain text files."""
        try:
            return path.read_text(errors="ignore")
        except Exception as e:
            raise ValueError(f"Failed to extract text from file: {str(e)}")
    
    @staticmethod
    def extract_emails(text: str) -> List[str]:
        """Extract email addresses from text."""
        pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        return list(set(re.findall(pattern, text)))
    
    @staticmethod
    def extract_phones(text: str) -> List[str]:
        """Extract phone numbers from text."""
        pattern = r"\b(?:\+?1[-.\s]?)?(?:\(?[0-9]{3}\)?[-.\s]?)?[0-9]{3}[-.\s]?[0-9]{4}\b"
        matches = re.findall(pattern, text)
        return [m.strip() for m in set(matches)]
    
    @staticmethod
    def extract_urls(text: str) -> List[str]:
        """Extract URLs from text."""
        pattern = r"https?://[^\s]+"
        return list(set(re.findall(pattern, text)))
    
    @staticmethod
    def extract_dates(text: str) -> List[str]:
        """Extract date patterns from text."""
        patterns = [
            r"\d{1,2}/\d{1,2}/\d{2,4}",  # MM/DD/YYYY
            r"\d{1,2}-\d{1,2}-\d{2,4}",  # MM-DD-YYYY
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b",
            r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December) \d{1,2},? \d{4}\b",
        ]
        dates = []
        for pattern in patterns:
            dates.extend(re.findall(pattern, text, re.IGNORECASE))
        return list(set(dates))
    
    @staticmethod
    def extract_skills(text: str) -> List[str]:
        """Extract skills from text using keyword matching."""
        text_lower = text.lower()
        skills = []
        
        for skill in MetadataExtractor.SKILL_KEYWORDS:
            # Use word boundaries to avoid partial matches
            pattern = r"\b" + re.escape(skill) + r"\b"
            if re.search(pattern, text_lower):
                skills.append(skill)
        
        return skills
    
    @staticmethod
    def extract_degrees(text: str) -> List[str]:
        """Extract educational degrees from text."""
        degrees = []
        text_lower = text.lower()
        
        for pattern in MetadataExtractor.DEGREE_PATTERNS:
            matches = re.findall(pattern, text_lower)
            degrees.extend(matches)
        
        return list(set(degrees))
    
    @staticmethod
    def extract_locations(text: str) -> List[str]:
        """Extract potential location mentions from text."""
        # Simple pattern for city, state/country combinations
        pattern = r"\b(?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*(?:[A-Z]{2}|[A-Za-z\s]+)\b"
        matches = re.findall(pattern, text)
        return list(set(matches))
    
    @staticmethod
    def detect_sections(text: str) -> Dict[str, str]:
        """Detect common resume/document sections."""
        sections = {}
        section_keywords = {
            "experience": r"(?:work experience|employment|professional experience|career history)",
            "education": r"(?:education|academic|qualifications|degree|university)",
            "skills": r"(?:skills|technical skills|core competencies|expertise)",
            "projects": r"(?:projects|portfolio|achievements|accomplishments)",
            "certifications": r"(?:certifications|licenses|credentials|certified)",
            "summary": r"(?:summary|objective|professional summary|about)",
            "contact": r"(?:contact|phone|email|address|linkedin)",
        }
        
        lines = text.split("\n")
        current_section = None
        section_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if this line matches a section header
            matched_section = None
            for section_name, pattern in section_keywords.items():
                if re.search(pattern, line_lower) and len(line) < 100:  # Likely a header
                    if current_section and section_content:
                        sections[current_section] = "\n".join(section_content).strip()
                    current_section = section_name
                    section_content = []
                    matched_section = section_name
                    break
            
            if not matched_section and current_section:
                section_content.append(line)
        
        if current_section and section_content:
            sections[current_section] = "\n".join(section_content).strip()
        
        return sections
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """
        Extract comprehensive metadata from a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            DocumentMetadata object with all extracted information
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        metadata = DocumentMetadata()
        
        # Extract basic file information
        file_meta = self._get_file_metadata(path)
        metadata.file_name = file_meta["file_name"]
        metadata.file_type = file_meta["file_type"]
        metadata.file_size = file_meta["file_size"]
        metadata.file_path = file_meta["file_path"]
        
        # Extract text based on file type
        text = ""
        suffix = path.suffix.lower()
        
        if suffix == ".pdf":
            text, has_selectable, page_count = self._extract_text_from_pdf(path)
            metadata.has_selectable_text = has_selectable
            metadata.page_count = page_count
            metadata.is_scanned = not has_selectable
            metadata.content_type = "scanned_document" if not has_selectable else "selectable_text"
        elif suffix in {".docx", ".doc"}:
            text, _ = self._extract_text_from_docx(path)
            metadata.page_count = None
            metadata.has_selectable_text = True
            metadata.content_type = "selectable_text"
        elif suffix in {".txt", ".md"}:
            text = self._extract_text_from_plain_text(path)
            metadata.has_selectable_text = True
            metadata.content_type = "selectable_text"
        else:
            # Treat as plain text
            text = self._extract_text_from_plain_text(path)
            metadata.has_selectable_text = True
            metadata.content_type = "selectable_text"
        
        # Set content metadata
        metadata.text_length = len(text)
        
        # Extract all entities
        metadata.entities["emails"] = self.extract_emails(text)
        metadata.entities["phones"] = self.extract_phones(text)
        metadata.entities["urls"] = self.extract_urls(text)
        metadata.entities["dates"] = self.extract_dates(text)
        metadata.entities["locations"] = self.extract_locations(text)
        metadata.entities["skills"] = self.extract_skills(text)
        metadata.entities["degrees"] = self.extract_degrees(text)
        
        # Detect sections
        metadata.sections = self.detect_sections(text)
        
        # Set MIME type based on file type
        mime_types = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
            ".md": "text/markdown",
        }
        metadata.content_type = mime_types.get(suffix, "application/octet-stream")
        
        return metadata
    
    def metadata_to_dict(self, metadata: DocumentMetadata) -> Dict[str, Any]:
        """Convert DocumentMetadata object to dictionary."""
        return {
            "file_name": metadata.file_name,
            "file_type": metadata.file_type,
            "file_size": metadata.file_size,
            "file_path": metadata.file_path,
            "extracted_at": metadata.extracted_at,
            "content_type": metadata.content_type,
            "page_count": metadata.page_count,
            "language": metadata.language,
            "has_selectable_text": metadata.has_selectable_text,
            "has_images": metadata.has_images,
            "image_count": metadata.image_count,
            "is_scanned": metadata.is_scanned,
            "ocr_confidence": metadata.ocr_confidence,
            "text_length": metadata.text_length,
            "sections": metadata.sections,
            "entities": metadata.entities,
            "metadata_tags": metadata.metadata_tags,
        }


# Convenience function
def extract_metadata(file_path: str) -> Dict[str, Any]:
    """
    Convenience function to extract metadata from a file.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Dictionary with extracted metadata
    """
    extractor = MetadataExtractor()
    metadata = extractor.extract_metadata(file_path)
    return extractor.metadata_to_dict(metadata)
